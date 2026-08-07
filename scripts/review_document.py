#!/usr/bin/env python3
"""读取已有 DOCX，输出格式审查和内容核验的结构化基础结果。

本脚本只负责安全读取 Word 结构和提取待核验文本，不调用搜索接口，也不改写原文件。
内容核验由 Agent 根据 claims 设计搜索方案并调用 dkag_search.py 完成。
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None


SKILL_ROOT = Path(__file__).resolve().parent.parent
INPUT_DIR = SKILL_ROOT / "official-docs" / "input"
OUTPUT_DIR = SKILL_ROOT / "official-docs" / "output"


def safe_path(value: str, allowed: Path, suffix: str) -> Path:
    raw = Path(value).expanduser()
    path = raw.resolve() if raw.is_absolute() else (allowed / raw.name).resolve()
    try:
        path.relative_to(allowed.resolve())
    except ValueError as exc:
        raise ValueError(f"文件必须位于 {allowed}: {value}") from exc
    if path.suffix.lower() != suffix:
        raise ValueError(f"文件必须是 {suffix} 文件: {value}")
    return path


def paragraph_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def inspect_document(path: Path) -> dict:
    if Document is None:
        raise RuntimeError("缺少 python-docx，请先按初始化流程安装依赖")
    doc = Document(str(path))
    paragraphs = []
    claims = []
    for index, paragraph in enumerate(doc.paragraphs, 1):
        text = paragraph_text(paragraph)
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        item = {"index": index, "text": text, "style": style}
        paragraphs.append(item)
        if re.search(r"政策|办法|条例|规定|标准|数据|达到|超过|占比|金额|亿元|万人|全国首个|唯一|领先|发布|印发|文号", text):
            claims.append({"paragraph": index, "text": text, "reason": "包含政策、数据、案例或高风险事实表述，建议通过深知可信搜索核验"})

    tables = []
    for table_index, table in enumerate(doc.tables, 1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append({"index": table_index, "rows": len(rows), "columns": max((len(row) for row in rows), default=0), "preview": rows[:3]})

    sections = []
    for index, section in enumerate(doc.sections, 1):
        sections.append({
            "index": index,
            "page_width_twips": section.page_width.twips if section.page_width else None,
            "page_height_twips": section.page_height.twips if section.page_height else None,
            "top_margin_twips": section.top_margin.twips if section.top_margin else None,
            "bottom_margin_twips": section.bottom_margin.twips if section.bottom_margin else None,
            "left_margin_twips": section.left_margin.twips if section.left_margin else None,
            "right_margin_twips": section.right_margin.twips if section.right_margin else None,
        })

    fonts = {}
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            name = run.font.name or "未明确设置"
            fonts[name] = fonts.get(name, 0) + len(run.text or "")

    return {
        "source_file": path.name,
        "format_review": {
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "sections": sections,
            "font_usage": fonts,
            "paragraphs": paragraphs,
            "tables": tables,
        },
        "content_review": {"claims": claims, "claim_count": len(claims)},
        "next_step": "仅格式审查可直接根据 format_review 输出问题；涉及内容真实性时，先根据 claims 设计搜索方案并经用户确认，再调用深知搜索。",
    }


def main() -> None:
    parser = argparse.ArgumentParser(description="读取已有 DOCX 并输出结构化审查基础结果")
    parser.add_argument("input", help="已有 DOCX，必须位于 official-docs/input")
    parser.add_argument("--output", "-o", help="JSON 输出文件，默认写入 official-docs/output")
    args = parser.parse_args()
    input_path = safe_path(args.input, INPUT_DIR, ".docx")
    if not input_path.exists():
        raise FileNotFoundError(f"输入文件不存在: {input_path}")
    output_path = safe_path(args.output, OUTPUT_DIR, ".json") if args.output else None
    result = inspect_document(input_path)
    rendered = json.dumps(result, ensure_ascii=False, indent=2)
    if output_path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(rendered, encoding="utf-8")
        print(f"审查基础结果已写入: {output_path}")
    else:
        print(rendered)


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"错误: {exc}", file=sys.stderr)
        raise SystemExit(1)
